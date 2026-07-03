# -*- coding: utf-8 -*-
"""
Generate Word (.docx) and PDF documents from the research report.
Font: 標楷體 (DFKai-SB), Size: 12pt
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import re

# ─── Configuration ───
FONT_NAME = '標楷體'
FONT_NAME_EN = 'DFKai-SB'
FONT_SIZE = Pt(12)
TITLE_SIZE = Pt(22)
SUBTITLE_SIZE = Pt(16)
H1_SIZE = Pt(18)
H2_SIZE = Pt(15)
H3_SIZE = Pt(13)
LINE_SPACING = 1.5
OUTPUT_DIR = r'c:\Users\coolexam\Documents\hackathon'

def set_run_font(run, size=FONT_SIZE, bold=False, color=None, italic=False):
    """Set font properties for a run."""
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Set Chinese font
    run.font.name = FONT_NAME
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         space_before=Pt(6), space_after=Pt(6),
                         line_spacing=LINE_SPACING, first_line_indent=None):
    """Set paragraph formatting."""
    para.alignment = alignment
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent

def add_heading_custom(doc, text, level=1):
    """Add a styled heading."""
    para = doc.add_paragraph()
    if level == 0:  # Main title
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER, Pt(24), Pt(12))
        run = para.add_run(text)
        set_run_font(run, TITLE_SIZE, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
    elif level == 1:  # Chapter title
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(18), Pt(8))
        run = para.add_run(text)
        set_run_font(run, H1_SIZE, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
        # Add bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C6E"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:  # Section title
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(14), Pt(6))
        run = para.add_run(text)
        set_run_font(run, H2_SIZE, bold=True, color=RGBColor(0x2D, 0x5F, 0x8A))
    elif level == 3:  # Subsection title
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(10), Pt(4))
        run = para.add_run(text)
        set_run_font(run, H3_SIZE, bold=True, color=RGBColor(0x3A, 0x7C, 0xA5))
    return para

def add_body_text(doc, text, indent=True):
    """Add body paragraph."""
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent=Cm(0.85) if indent else None)
    run = para.add_run(text)
    set_run_font(run)
    return para

def add_body_rich(doc, segments, indent=True):
    """Add body paragraph with mixed formatting.
    segments: list of (text, bold, italic, color) tuples
    """
    para = doc.add_paragraph()
    set_paragraph_format(para, first_line_indent=Cm(0.85) if indent else None)
    for text, bold, italic, color in segments:
        run = para.add_run(text)
        set_run_font(run, bold=bold, italic=italic, color=color)
    return para

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph()
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(2), Pt(2))
    para.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    para.paragraph_format.first_line_indent = Cm(-0.5)
    # Parse bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run('● ' if '● ' not in text else '')
            set_run_font(run)
            run2 = para.add_run(part[2:-2])
            set_run_font(run2, bold=True)
        else:
            if parts.index(part) == 0 and not part.startswith('●'):
                run = para.add_run('● ' + part)
            else:
                run = para.add_run(part)
            set_run_font(run)
    return para

def add_quote(doc, text):
    """Add a quote block."""
    para = doc.add_paragraph()
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(8), Pt(8))
    para.paragraph_format.left_indent = Cm(1.5)
    para.paragraph_format.right_indent = Cm(1.5)
    # Add left border
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:left w:val="single" w:sz="18" w:space="8" w:color="3A7CA5"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    # Add shading
    shading = parse_xml(
        '<w:shd {} w:fill="EBF5FB" w:val="clear"/>'.format(nsdecls('w'))
    )
    pPr.append(shading)
    run = para.add_run(text)
    set_run_font(run, italic=True, color=RGBColor(0x2D, 0x5F, 0x8A))
    return para

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        set_run_font(run, Pt(11), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Header background
        shading = parse_xml(
            '<w:shd {} w:fill="1A3C6E" w:val="clear"/>'.format(nsdecls('w'))
        )
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Style data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(cell_text))
            set_run_font(run, Pt(11))
            # Alternating row colors
            if r_idx % 2 == 0:
                shading = parse_xml(
                    '<w:shd {} w:fill="F2F7FB" w:val="clear"/>'.format(nsdecls('w'))
                )
                cell._tc.get_or_add_tcPr().append(shading)
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    # Add table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    borders = parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # Space after table
    doc.add_paragraph()
    return table

def add_separator(doc):
    """Add a horizontal line separator."""
    para = doc.add_paragraph()
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(12))
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

def add_dialogue_line(doc, speaker, text, is_ai=False):
    """Add a dialogue line for the conversation example."""
    para = doc.add_paragraph()
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(3), Pt(3))
    para.paragraph_format.left_indent = Cm(2.0)
    
    run_speaker = para.add_run(f'{speaker}：')
    set_run_font(run_speaker, bold=True, 
                 color=RGBColor(0x3A, 0x7C, 0xA5) if is_ai else RGBColor(0x8B, 0x45, 0x13))
    
    run_text = para.add_run(text)
    set_run_font(run_text, italic=True if is_ai else False)

def build_document():
    """Build the complete Word document."""
    doc = Document()
    
    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # ─── Set default font ───
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.line_spacing = LINE_SPACING
    
    # ═══════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()
    
    add_heading_custom(doc, '以人本思維重構臺灣城鄉教育落差', 0)
    add_heading_custom(doc, '之普惠科技應用研究', 0)
    
    doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    set_paragraph_format(subtitle_para, WD_ALIGN_PARAGRAPH.CENTER, Pt(12), Pt(6))
    run = subtitle_para.add_run('——「說說學伴」')
    set_run_font(run, SUBTITLE_SIZE, color=RGBColor(0x3A, 0x7C, 0xA5))
    
    subtitle_para2 = doc.add_paragraph()
    set_paragraph_format(subtitle_para2, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(6))
    run2 = subtitle_para2.add_run('實體智能無螢幕伴讀系統方案')
    set_run_font(run2, SUBTITLE_SIZE, color=RGBColor(0x3A, 0x7C, 0xA5))
    
    for _ in range(2):
        doc.add_paragraph()
    
    info_para = doc.add_paragraph()
    set_paragraph_format(info_para, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(6))
    run = info_para.add_run('2026 雲湧智生：臺灣生成式 AI 應用黑客松競賽')
    set_run_font(run, Pt(13), color=RGBColor(0x66, 0x66, 0x66))
    
    info_para2 = doc.add_paragraph()
    set_paragraph_format(info_para2, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(6))
    run = info_para2.add_run('創意交流組——晶創未來')
    set_run_font(run, Pt(13), color=RGBColor(0x66, 0x66, 0x66))
    
    doc.add_paragraph()
    
    team_para = doc.add_paragraph()
    set_paragraph_format(team_para, WD_ALIGN_PARAGRAPH.CENTER, Pt(6), Pt(6))
    run = team_para.add_run('師大酷英團隊')
    set_run_font(run, Pt(13), color=RGBColor(0x66, 0x66, 0x66), bold=True)
    
    members_para = doc.add_paragraph()
    set_paragraph_format(members_para, WD_ALIGN_PARAGRAPH.LEFT, Pt(6), Pt(6))
    members_para.paragraph_format.left_indent = Cm(6.4)
    
    run_label = members_para.add_run('成員：')
    set_run_font(run_label, Pt(13), color=RGBColor(0x66, 0x66, 0x66))
    
    run_names = members_para.add_run(' 郭泰源\n　　　 王智城\n　　　 李宗哲\n　　　 范鈞淯')
    set_run_font(run_names, Pt(13), color=RGBColor(0x66, 0x66, 0x66))
    
    # Page break after cover
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════
    # TABLE OF CONTENTS (Manual)
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '目　錄', 0)
    doc.add_paragraph()
    
    toc_items = [
        ('摘要', ''),
        ('第一章　研究背景與問題意識', ''),
        ('　1.1　臺灣社會變遷與教育平等的迫切性', ''),
        ('　1.2　普惠科技的人本精神', ''),
        ('第二章　痛點深度剖析：偏鄉雙語教育的三重困境', ''),
        ('　2.1　困境一：雙語政策下的「雙峰現象」', ''),
        ('　2.2　困境二：學習低成就學生的「習得性無助」', ''),
        ('　2.3　困境三：數位教學工具的瓶頸與教師行政負擔', ''),
        ('　2.4　痛點總結：為什麼需要一個全新的解法？', ''),
        ('第三章　解決方案：「說說學伴」系統設計', ''),
        ('　3.1　方案定位與核心理念', ''),
        ('　3.2　系統架構：邊緣與雲端的雙網混成設計', ''),
        ('　3.3　鷹架式語伴教學法', ''),
        ('　3.4　針對 8-12 歲學童之主動引導與防分心沉迷設計', ''),
        ('　3.5　科技酷玩：符合 8-12 歲學童偏好之外觀美學', ''),
        ('第四章　兩大普惠應用情境設計', ''),
        ('　4.1　情境一：偏鄉小學學習扶助課堂的「無壓力伴讀鷹架」', ''),
        ('　4.2　情境二：偏鄉社區跨世代教育共創——「在地記憶雙語繪本」', ''),
        ('第五章　國產晶片與邊緣 AI 運算的使能價值', ''),
        ('　5.1　為什麼需要邊緣 AI？', ''),
        ('　5.2　國產晶片開發平台之技術規格', ''),
        ('　5.3　邊緣 AI 推論效能與離線能力升級', ''),
        ('第六章　社會影響力與未來展望', ''),
        ('　6.1　預期社會影響力', ''),
        ('　6.2　落地路徑規劃', ''),
        ('　6.3　未來擴展藍圖', ''),
        ('結論', ''),
        ('參考文獻', ''),
    ]
    
    for item, _ in toc_items:
        para = doc.add_paragraph()
        is_chapter = not item.startswith('　')
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(2), Pt(2))
        if is_chapter:
            para.paragraph_format.space_before = Pt(6)
        run = para.add_run(item)
        set_run_font(run, Pt(12), bold=is_chapter)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════
    # ABSTRACT
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '摘　要', 1)
    
    add_body_text(doc, 
        '臺灣社會正面臨前所未有的結構性轉變，其中城鄉教育資源分配不均所導致的「數位落差」與「雙語政策焦慮」，'
        '已成為下一個十年最迫切的社會變遷議題之一。本研究以「永續城鄉與教育平等」為唯一聚焦痛點，深入剖析偏鄉'
        '國小 3 至 6 年級學習扶助學童在雙語教育推動過程中所面臨的「習得性無助」與「數位鴻溝」困境，並以人本思維'
        '為核心，結合國產晶片邊緣運算技術與 AWS 雲端生成式 AI，提出一套名為「說說學伴」的實體'
        '智能無螢幕伴讀系統方案。')
    
    add_body_text(doc,
        '本方案以「無螢幕實體伴讀裝置」為核心創新，搭配「雙網混成推論架構」與「鷹架式語伴教學法」，在保障隱私、'
        '克服偏鄉斷網限制的前提下，為弱勢學童提供有溫度的 AI 英語學習陪伴，同時減輕基層教師的行政負擔，真正實踐'
        '以科技促進教育平權的普惠價值。')
    
    add_separator(doc)
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 1
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '第一章　研究背景與問題意識', 1)
    
    add_heading_custom(doc, '1.1　臺灣社會變遷與教育平等的迫切性', 2)
    
    add_body_text(doc,
        '臺灣社會正經歷多重結構性變遷——超高齡化社會的提前到來、勞動力結構失衡、環境災害風險增加，以及城鄉資源'
        '分配不均。在眾多社會議題中，「教育平等」是最具根源性的課題：教育是社會流動的基石，若城鄉間的教育品質持續'
        '失衡，將進一步加劇社會階層固化，削弱國家整體競爭力與社會韌性。')
    
    add_body_text(doc,
        '「2026 雲湧智生：臺灣生成式 AI 應用黑客松競賽」由臺灣半導體產學研發聯盟（TIARA）等單位聯合發起，其'
        '「晶創未來」命題精神在於引導跨領域團隊，運用人文社會思維洞察真實需求，結合生成式 AI 技術與電子數位科技，'
        '提出具備社會影響力且能實際落地的普惠科技解決方案。本研究即以此精神為指引，聚焦「永續城鄉與教育平等」'
        '單一痛點，深度回應臺灣教育現場最真實的困境。')
    
    add_heading_custom(doc, '1.2　普惠科技的人本精神', 2)
    
    add_body_text(doc,
        '普惠科技（Inclusive Technology）的核心價值在於「人本思維」——科技的設計不應僅追求技術指標的突破，而應'
        '優先考量弱勢群體、偏鄉住民及學習低成就學生的使用近便性，消除數位排斥，使技術進步轉化為推動社會公平與'
        '教育永續發展的助力。')
    
    add_body_text(doc, '在本研究中，「人本思維」具體體現為三個設計原則：', indent=False)
    
    add_bullet(doc, '**降低使用門檻**：讓不熟悉數位裝置的偏鄉學童「說話即用」，無需操作螢幕選單。')
    add_bullet(doc, '**消除學習焦慮**：以溫暖的實體陪伴取代冰冷的測驗介面，重建學童的學習自信心。')
    add_bullet(doc, '**減輕教師負擔**：透過 AI 自動生成教學分析報表，讓教師專注於教學本身。')
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 2
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '第二章　痛點深度剖析：偏鄉雙語教育的三重困境', 1)
    
    add_heading_custom(doc, '2.1　困境一：雙語政策下的「雙峰現象」', 2)
    
    add_body_text(doc,
        '臺灣自推動雙語教育與「2030 雙語政策」以來，各級學校紛紛加速推動學科全英語授課（EMI）及雙語化教學。'
        '然而，基層教學現場卻因此產生極大張力。')
    
    add_body_text(doc,
        '由於城鄉之間數位設備與英語陪伴資源的嚴重失衡，偏鄉及經濟弱勢學生因缺乏家庭支持系統，在面臨全英語或'
        '雙語學科教學時極易產生挫折感。相較於都會區學生擁有豐富的課後英語補習資源與家庭陪伴，偏鄉學童往往在'
        '課堂結束後便失去所有英語接觸的機會，導致學習成就的雙峰化現象日益惡化，甚至引發基層教師與家長團體的'
        '強烈反彈。')

    add_body_text(doc, '關鍵數據佐證：', indent=False)
    add_bullet(doc, '**學力落差巨大**：在英語學習扶助篩選測驗中，偏遠地區國中小學童的未通過率為全國平均的 1.3 至 2.4 倍。在國中教育會考中，偏鄉英語達「待加強 (C)」的比例長期高達 47% 以上（都會區僅為 23.66% - 32.25%，差距達 1.78 至 1.90 倍）。以花蓮縣為例，偏鄉英語待加強比例高達 74.15%，是同縣非偏鄉的 1.75 倍。')
    add_bullet(doc, '**師資流動率與短缺**：偏鄉學校常年面臨招聘困難，代理教師比例高達 20%，常態性因流動率過高導致學童「每學期都在換英語老師」，且因招考多次無人應聘，普遍存在英語非專長授課（兼任）的斷層問題。')
    add_bullet(doc, '**硬體齊備但應用真空**：政府在 2022-2025 年推動「生生用平板」方案，累計投入 200 億元、配發 61 萬台平板，偏鄉學校已達 1:1（一人一機）的高普及率。然而，受限於斷網痛點，缺乏能在本地運作、支援流暢口說互動的 AI 應用，使大量平板流於播放影片或靜態紙筆化測驗，形成嚴重的「軟體與應用真空」。')
    
    add_heading_custom(doc, '2.2　困境二：學習低成就學生的「習得性無助」', 2)
    
    add_body_text(doc,
        '在教育部積極推動的「學習扶助（補救教學）」機制中，低成就學生常因自我效能感低落，對傳統紙筆測驗或生硬的'
        '線上評量介面產生焦慮，進入抗拒學習的惡性循環。對於國小 3 至 6 年級（8-12 歲）的偏鄉學扶學童而言，'
        '這種「習得性無助」尤為嚴重：')
    
    add_bullet(doc, '**不敢開口**：害怕發音不標準被同學嘲笑，在課堂上選擇沉默。')
    add_bullet(doc, '**抗拒螢幕**：傳統數位學習平台的介面設計類似考試，引發焦慮而非興趣。')
    add_bullet(doc, '**缺乏陪伴**：放學後無人陪伴練習英語口說，學習動機在孤獨中消磨殆盡。')
    
    add_heading_custom(doc, '2.3　困境三：數位教學工具的瓶頸與教師行政負擔', 2)
    
    add_body_text(doc,
        '教育部委託國立臺灣師範大學建置「Cool English 英語線上學習平台」（酷英網），已累積超過 270 萬註冊用戶，'
        '並於 114 年 6 月推出大專專區。平台已具備 AI 功能，如 CoolEBot 聊天機器人與語音辨識系統，是目前臺灣'
        '最具規模的英語數位學習資源。然而，現有工具仍存在以下瓶頸：')
    
    add_table(doc,
        ['瓶頸面向', '具體問題', '對偏鄉學童的影響'],
        [
            ['網路依賴', '所有 AI 功能需穩定網路連線', '偏鄉斷網時工具完全失效'],
            ['螢幕介面', '純數位互動缺乏物理實體陪伴', '學童注意力易渙散，缺乏安全感'],
            ['教師負擔', '派課、追蹤進度、分析報表皆需手動', '行政時間擠壓教學時間'],
            ['互動模式', '以文字輸入為主，口說功能有限', '無法解決「不敢開口」的核心問題'],
        ],
        [3.5, 5.5, 5.5]
    )
    
    add_heading_custom(doc, '2.4　痛點總結：為什麼需要一個全新的解法？', 2)
    
    add_body_text(doc, '綜合以上三重困境，臺灣偏鄉雙語教育的核心矛盾可歸結為：', indent=False)
    
    add_quote(doc, '政策要求學生「開口說英語」，但現有的數位工具無法讓偏鄉學童「安心地開口」。')
    
    add_body_text(doc,
        '這個矛盾無法僅靠改良現有平台來解決，而需要一個從「硬體形式」到「互動模式」到「教學法」全面重新思考'
        '的創新方案——這正是「說說學伴」誕生的原點。')
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 3
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '第三章　解決方案：「說說學伴」系統設計', 1)
    
    add_heading_custom(doc, '3.1　方案定位與核心理念', 2)
    
    add_body_text(doc,
        '「說說學伴」不是一個 App，不是一個網頁平台，而是一個有溫度的物理終端——一個偏鄉學童可以抱在懷中、'
        '用說話就能互動的 AI 英語學習夥伴。其核心理念可用三個「不」來概括：')
    
    add_bullet(doc, '**不需要螢幕**：消除螢幕帶來的注意力分散與考試焦慮。')
    add_bullet(doc, '**不需要網路也能用**：邊緣 AI 確保偏鄉斷網時基礎功能不中斷。')
    add_bullet(doc, '**不需要會操作手機**：說話即用，徹底消除數位鴻溝。')
    
    add_heading_custom(doc, '為什麼不是 App？——與傳統方案的結構性差異', 3)
    
    add_table(doc,
        ['比較維度', '傳統 App / 網頁平台', '說說學伴'],
        [
            ['螢幕分散', '✘ 學童易被其他 App 分心', '✓ 無螢幕，專注語音互動'],
            ['學習焦慮', '✘ 介面像考試，引發壓力', '✓ 互動像聊天，輕鬆自然'],
            ['斷網可用', '✘ 完全失效', '✓ 邊緣 AI 即時接手'],
            ['數位鴻溝', '✘ 需會操作觸控介面', '✓ 說話即用，零門檻'],
            ['情感連結', '✘ 冰冷的螢幕', '✓ 可擁抱的實體夥伴'],
        ],
        [3.0, 5.5, 5.5]
    )
    
    add_heading_custom(doc, '3.2　系統架構：邊緣與雲端的雙網混成設計', 2)
    
    add_body_text(doc,
        '「說說學伴」採用「雙網混成推論」(Hybrid Edge-Cloud Inference) 架構，根據網路環境自動切換運算模式，'
        '確保在任何情境下都能為學童提供不中斷的學習陪伴。')
    
    add_heading_custom(doc, '3.2.1　本地邊緣端：精準控制與隱私安全', 3)
    
    add_bullet(doc, '搭載聯發科 Genio 520 處理器，利用 10 TOPS NPU 執行高效能 Breeze-ASR（自動語音辨識）與 KWS（關鍵字喚醒）。')
    add_bullet(doc, '學生可透過離線語音指令直接喚醒系統（如「說說學伴，我想讀故事」），無須連接網路即可進行即時的基礎發音評測與跟讀反饋。')
    add_bullet(doc, '所有語音資料於晶片內即時處理完畢，確保「不落地、不升雲」的最高規格隱私保障。')
    
    add_heading_custom(doc, '3.2.2　雲端生成層：AWS Bedrock 彈性架構', 3)
    
    add_body_text(doc,
        '當系統偵測到環境 Wi-Fi 時，自動啟動「雙網混成模式」，將學生的進階對話語音特徵傳輸至 AWS 雲端。')
    
    add_body_text(doc,
        '檢索增強生成 (RAG) 技術：系統後台串接多個臺灣本土開放教育資源，確保 AI 生成內容完全契合臺灣課綱體制：')
    
    add_bullet(doc, '**Cool English 酷英網**：教育部官方英語學習平台，涵蓋國中小、高中職及大專之聽力、字彙、口說技巧與歷屆試題。')
    add_bullet(doc, '**均一教育平台**：臺灣最大公益教育平台，提供跨學科學習資源。')
    add_bullet(doc, '**國教院教學媒體**：國家教育研究院之官方教學素材與課綱對照資源。')
    
    add_body_text(doc,
        '反思教學 AI 助教：利用大語言模型對學生的口說內容進行句型與用字偵錯，自動為教師生成班級學習進度與'
        '個別學生的學習弱點報表，協助教師進行「差異化教學」，免去繁瑣的手動派課與數據解讀工作。')
    
    add_heading_custom(doc, '3.3　鷹架式語伴教學法', 2)
    
    add_body_text(doc,
        '「說說學伴」的互動設計採用語言教育學中的「鷹架引導 (Scaffolding)」理論，其最大特色是允許中英混合輸入，'
        '以溫和的語氣逐步引導學童從母語過渡至英語表達。')
    
    add_body_text(doc, '【對話範例：偏鄉學扶學童「阿明」的互動歷程】', indent=False)
    doc.add_paragraph()
    add_dialogue_line(doc, '阿明', '「我想吃 apple。」', is_ai=False)
    add_dialogue_line(doc, '說說學伴', '「哇，阿明想吃蘋果！跟著我說一遍：I want to eat an apple.」', is_ai=True)
    add_dialogue_line(doc, '阿明', '「I want to eat... an apple.」', is_ai=False)
    add_dialogue_line(doc, '說說學伴', '「太棒了！你說得很好！apple 的 /æ/ 音再大聲一點會更好聽喔。我們再來一次好嗎？」', is_ai=True)
    doc.add_paragraph()
    
    add_body_text(doc, '這種互動模式的關鍵設計原則：', indent=False)
    add_bullet(doc, '**不糾正，而是引導**：AI 不直接指出錯誤，而是用正面語氣示範正確用法。')
    add_bullet(doc, '**允許混合語言**：不要求學童一開始就用全英語，降低開口的心理障礙。')
    add_bullet(doc, '**即時正面回饋**：每次嘗試都給予鼓勵，重建學習自信心。')
    add_bullet(doc, '**多元語音選擇**：提供不同腔調、人聲與語速選擇，適應不同學習階段。')

    add_heading_custom(doc, '3.4　針對 8-12 歲學童之主動引導與防分心沉迷設計', 2)
    add_body_text(doc, '針對中高年級國小學童（8-12 歲）的心理與學習特性，「說說學伴」在互動機制上引入了主動化、遊戲化與防沉迷的設計：')
    add_bullet(doc, '**主動式聊天與興趣導向引導 (Proactive Engagement)**：系統能主動依據時間或學習紀錄啟動話題，並內建學童感興趣的流行音樂、熱門遊戲（如 Minecraft, Roblox）、經典動漫與電影角色等主題語庫。')
    add_bullet(doc, '**遊戲化挑戰機制 (Fast, Fun, and Fair)**：設計 2-3 分鐘的短時效任務提供即時回饋，每次跟讀成功會播放專屬音效並解鎖新章節，藉由邊緣/雲端評測給予透明公正的發音星等評級，引導學童反覆挑戰。')
    add_bullet(doc, '**擬人化防沉迷休息機制 (Time-Out Mechanism)**：設定單次使用時限（如 25-30 分鐘，對齊番茄工作法），超時後 AI 會以擬人化方式宣告需要睡覺休息並自動休眠，建立良好科技使用習慣。')

    add_heading_custom(doc, '3.5　科技酷玩：符合 8-12 歲學童偏好之外觀美學', 2)
    add_body_text(doc, '為了讓這個無螢幕裝置對 8-12 歲學童產生持久的吸引力，外觀設計摒棄了「低齡化」的玩具感，走向「精緻酷玩」路線：')
    add_bullet(doc, '**「科技酷玩 (Tech-Toy / Maker)」視覺風格**：外型結合簡約幾何線條，採用石墨灰、極光白等低飽和度科技配色與防撞邊條，呈現高端電子或智慧潮玩質感，並提供個性化磁吸裝飾與客製化貼紙。')
    add_bullet(doc, '**溫暖手感的雙材質拼接**：採用天然木質底座與親膚防摔霧面矽膠機身主體，讓學童捧在手中或抱在懷裡時感到安全放鬆。')
    add_bullet(doc, '**情緒動態 LED 呼吸光環 (Emotional Light Ring)**：設備頂端配置多色彩 LED 動態光環。呼吸燈光能展現「情緒生命感」：淺綠色呼吸代表 AI 正在聆聽，暖黃色代表思考，炫彩漸變代表挑戰成功，以光效代替冰冷螢幕。')
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 4
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '第四章　兩大普惠應用情境設計', 1)
    
    add_heading_custom(doc, '4.1　情境一：偏鄉小學學習扶助課堂的「無壓力伴讀鷹架」', 2)
    
    add_body_rich(doc, [
        ('場景描述：', True, False, RGBColor(0x1A, 0x3C, 0x6E)),
        ('在偏鄉補救教學課堂中，國小四年級的阿明常因自卑而不敢在全英語課堂上開口。其他同學已經能簡單對話，'
         '但阿明連最基本的自我介紹都說不出來。傳統的線上學習平台讓他覺得「又要考試了」，每次看到螢幕上的題目'
         '就緊張到手心冒汗。', False, False, None),
    ])
    
    add_body_rich(doc, [
        ('引入「說說學伴」後的改變：', True, False, RGBColor(0x1A, 0x3C, 0x6E)),
        ('阿明將木質外觀的「說說學伴」抱在懷中，透過實體喇叭進行溫馨的「雙語混成對話」。沒有螢幕、沒有分數、'
         '沒有紅色的錯誤標記——只有一個像朋友一樣耐心的聲音，陪他一句一句地練習。在離線模式下，邊緣晶片進行'
         '即時的基礎發音評估，阿明不用擔心網路延遲帶來的卡頓與挫折。當教室有 Wi-Fi 時，系統自動切換至雲端模式，'
         '提供更豐富的對話情境與進階語法引導。', False, False, None),
    ])
    
    add_body_rich(doc, [
        ('教師端價值：', True, False, RGBColor(0x1A, 0x3C, 0x6E)),
        ('導師林老師不再需要花費大量時間手動分析每位學扶學生的學習進度。「說說學伴」的 AI 助教自動生成一份清晰的'
         '班級報表，標示出阿明在「母音發音」與「簡單句型」上需要額外加強，讓林老師可以精準地進行差異化教學。',
         False, False, None),
    ])
    
    add_heading_custom(doc, '4.2　情境二：偏鄉社區跨世代教育共創——「在地記憶雙語繪本」', 2)
    
    add_body_text(doc,
        '在偏鄉社區關懷據點，國小五年級的小雅與社區的美濃客家阿嬤一起參加「跨世代教育共創」活動。互動流程如下：')
    
    add_bullet(doc, '**阿嬤口述**：阿嬤對著「說說學伴」以客語和國語講述美濃菸樓的歷史故事與年輕時的生活記憶。學伴本地端晶片進行語音採集與噪音過濾。')
    add_bullet(doc, '**AI 翻譯重構**：語音上傳至 AWS 雲端後，AI 將阿嬤的口述語音翻譯並重構為生動的雙語（中英）故事大綱。')
    add_bullet(doc, '**學童英語改寫**：小雅在學伴的鷹架引導下，嘗試用簡單的英語句子改寫阿嬤的故事，「說說學伴」即時提供語句修正與詞彙建議。')
    add_bullet(doc, '**繪本共創**：最終產出一本專屬的「AI 雙語有聲繪本」，記錄阿嬤的在地記憶，也記錄小雅的英語學習歷程。')
    
    add_body_text(doc, '此情境的雙重教育價值：', indent=False)
    add_bullet(doc, '**對學童**：學習動機從「為了考試」轉變為「為了記錄阿嬤的故事」，英語成為有溫度的溝通工具。')
    add_bullet(doc, '**對長者**：透過「敘事療癒」延緩認知退化，在地文化得以被記錄與傳承。')
    add_bullet(doc, '**對社區**：跨世代連結在科技媒介下被強化，展現「永續城鄉」的深層意涵。')
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 5
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '第五章　國產晶片與邊緣 AI 運算的使能價值', 1)
    
    add_heading_custom(doc, '5.1　為什麼需要邊緣 AI？', 2)
    
    add_body_text(doc,
        '在偏鄉、弱勢家庭及極端情境下，高頻寬、高成本的雲端 AI 往往因「網路中斷」或「資費高昂」而無法普及。'
        '「邊緣 AI（Edge AI）」是消除此類數位排除的關鍵。對於「說說學伴」的目標場景，邊緣 AI 的價值體現在三個層面：')
    
    add_bullet(doc, '**隱私保護 (Privacy by Design)**：未成年兒童的語音資料於晶片內即時處理完畢，不上傳雲端，提供最高規格隱私保障。')
    add_bullet(doc, '**擺脫網路限制 (Zero Network Dependency)**：偏遠山區或網路品質不穩的教室中，仍可提供即時語音互動服務。')
    add_bullet(doc, '**極致低耗能與低建置成本**：以電池供電可維持長時間運行，硬體造價低，利於政府在偏鄉大規模配發。')
    
    add_heading_custom(doc, '5.2　國產晶片開發平台之技術規格', 2)
    
    add_table(doc,
        ['認定類別', '核心定義', '本方案之對應'],
        [
            ['自主研發晶片', '核心 IP 由我國依法登記之企業自主研發，擁有晶片完整智慧財產權。',
             '聯發科技 (MediaTek)：Genio 520 應用處理器，內建 8 核心 CPU 與 10 TOPS AI 處理單元（NPU），聯發科為我國依法登記企業，擁有 Genio 晶片完整 IP。'],
            ['AI 開發生態系', '提供完整的 AI 軟體開發工具鏈與模型部署框架。',
             'NeuroPilot 8 AI SDK：支援 TensorFlow Lite、ONNX Runtime 等主流框架，可直接在 NPU 上加速 ASR / LLM / TTS 模型推論。'],
            ['開發板取得管道', '透過官方認可管道取得開發硬體。',
             'ideas Hatch 智造工具包：HUB G520 開發板可透過 ideas Hatch 平台申請取得，降低偏鄉教育場域的硬體導入門檻。'],
        ],
        [3.0, 5.0, 6.5]
    )
    
    add_heading_custom(doc, '5.3　邊緣 AI 推論效能與離線能力升級', 2)
    
    add_body_text(doc,
        '以聯發科 Genio 520 平台為核心，搭配 NeuroPilot 8 AI SDK 進行模型量化與 NPU 加速部署，實現完整的離線 AI 對話管線：')
    
    add_bullet(doc, 'NPU 算力：10 TOPS（INT8），足以同時運行 ASR、LLM、TTS 三大模型。')
    add_bullet(doc, '離線 ASR（Breeze-ASR / Sherpa-ONNX SenseVoice）：針對臺灣華語與英語進行即時語音辨識，延遲低於 200ms，支援 KWS 喚醒與連續語音輸入。')
    add_bullet(doc, '離線 LLM（Qwen3 1.7B via Ollama）：透過 INT4 量化部署於 NPU，推論速度約 20 tok/s，可即時生成鷹架式英語引導對話。')
    add_bullet(doc, '離線 TTS（Piper TTS / BreezyVoice）：低延遲語音合成，支援中英雙語自然朗讀，首音延遲低於 150ms。')
    add_body_text(doc,
        '此技術規格使「說說學伴」具備強大的離線對話能力，'
        '即便在完全斷網的偏鄉環境中，學童也能獲得流暢的 AI 伴讀互動體驗。')
    
    # ═══════════════════════════════════════════════════
    # CHAPTER 6
    add_heading_custom(doc, '第六章　社會影響力與未來展望', 1)
    
    add_heading_custom(doc, '6.1　預期社會影響力', 2)
    
    add_bullet(doc, '**降低偏鄉學童英語焦慮**：透過無壓力的伴讀互動，讓學習扶助成為「有朋友陪伴」的溫暖體驗。')
    add_bullet(doc, '**提升學習扶助通過率**：結合課綱對齊的 RAG 內容與個人化學習路徑，有效提升偏鄉學童英語基礎能力。')
    add_bullet(doc, '**減輕教師行政負擔**：AI 助教自動生成教學報表，將教師從繁瑣的數據分析中解放，回歸教學本質。')
    add_bullet(doc, '**彌平城鄉教育數位落差**：以低成本、零門檻的實體裝置，讓偏鄉學童享有同等品質的 AI 學習資源。')
    add_bullet(doc, '**傳承在地文化**：跨世代共創情境讓在地記憶被保存、翻譯、傳承，展現教育的文化永續價值。')
    
    add_heading_custom(doc, '6.2　落地路徑規劃', 2)
    
    add_bullet(doc, '**第一階段**：結合教育部「中小學數位學習精進方案」，在 5-10 所偏鄉國小進行實驗導入。')
    add_bullet(doc, '**第二階段**：與地方教育局、公益團體合作，以低硬體成本大規模配發至偏鄉學扶班級與社區關懷據點。')
    add_bullet(doc, '**第三階段**：開放邊緣-雲端混成架構為平台化服務，讓其他教育內容開發者可開發更多學科學習模組。')
    
    add_heading_custom(doc, '6.3　未來擴展藍圖', 2)
    
    add_body_text(doc,
        '「說說學伴」的邊緣-雲端混成架構具備高度可擴展性。在教育場景驗證成功後，同樣的技術核心可延伸至'
        '臺灣社會其他變遷議題：')
    
    add_bullet(doc, '**身心健康與弱勢關懷**：為獨居高齡長者提供免聯網的聲控陪伴與認知訓練裝置。')
    add_bullet(doc, '**人力結構與人才發展**：為技職體系學生提供英語口說模擬面試練習。')
    add_bullet(doc, '**環境安全與社會韌性**：於災害斷網情境下提供離線語音求救與安全指引系統。')
    
    # ═══════════════════════════════════════════════════
    # CONCLUSION
    # ═══════════════════════════════════════════════════
    add_heading_custom(doc, '結　論', 1)
    
    add_body_text(doc,
        '面對臺灣下一個十年的社會結構變遷，科技的發展不應只是追求冷冰冰的算力與商業獲利，而應回歸人本初心，'
        '去溫暖、去關懷那些在數位浪潮中被遺忘的群體。')
    
    add_quote(doc, '每個孩子都值得有人陪他開口。')
    
    add_body_text(doc,
        '無論是在偏鄉山區的教室裡、社區關懷據點的客廳中，還是在任何一個沒有網路的角落，「說說學伴」這個有溫度的'
        '小夥伴都能陪伴學童勇敢地說出第一句英語。')
    
    add_body_text(doc,
        '本研究展示了如何將我國最具優勢的半導體科技（國產聯發科 Genio 520 晶片平台）與最前沿的生成式 AI（AWS Bedrock '
        '雲端基礎模型）進行深度軟硬整合，在保障個人隱私、克服城鄉斷網限制的前提下，為偏鄉教育提供具備高度可行性'
        '與原創性的解答。')
    
    add_body_text(doc,
        '讓電子數位科技真正成為推動臺灣教育平權與社會永續發展的關鍵助力——這就是「晶創未來」的真正意義。')
    
    # ═══════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_custom(doc, '參考文獻', 1)
    
    references = [
        '創意交流組 命題文件 - 雲湧智生：臺灣生成式 AI 應用黑客松競賽.pdf',
        '2026 雲湧智生：臺灣生成式 AI 應用黑客松 - 比賽 - 獎金獵人. https://bhuntr.com/tw/competitions/mk3yl7w2q9r402gsb2',
        '雙語教育的優點與挑戰 - 親子天下 Shopping. https://shopping.parenting.com.tw/Article/Detail/99144',
        '2026 年「台美教育倡議」事實清單 - 美國在台協會. https://www.ait.org.tw/',
        '「雙語政策」有違國家語言發展法精神 - 公共政策網路參與平臺. https://join.gov.tw/',
        'Cool English 平臺 - 運用於英文學習扶助課程. https://www.ttbf.org.tw/',
        '關於酷英 | COOLENGLISH. https://www.coolenglish.edu.tw/',
        '英語系建置酷英平台結合 AI 技術 - 國立臺灣師範大學. https://pr.ntnu.edu.tw/',
        '酷英 (Cool English) 介紹 - 中山醫學大學. https://emi.csmu.edu.tw/',
        '嵌入式 AI 語音聲控邊緣運算晶片裝置 - 台灣創新技術博覽會. https://tie.twtm.com.tw/',
        '聯發科技 Genio 智慧物聯網平台 - 聯發科技. https://www.mediatek.tw/products/iot/genio-520',
        '聯發科技 NeuroPilot 開放生態系. https://neuropilot.mediatek.com/',
        'ideas Hatch 智造工具包與開發資源 - 財團法人資訊工業策進會. https://www.ideas-hatch.com/',
        '瑞昱 Ameba 開發板 - 超圖解系列圖書. https://swf.com.tw/',
        '教育部投入 1.51 億推動 271 所高中職雙語校園 - 自由時報. https://news.ltn.com.tw/',
        '推動中小學數位學習精進方案 - 臺北市教育局. https://www-ws.gov.taipei/',
    ]
    
    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph()
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, Pt(2), Pt(2))
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(-1.0)
        run = para.add_run(f'[{i}]　{ref}')
        set_run_font(run, Pt(10))
    
    # ─── Save ───
    docx_path = os.path.join(OUTPUT_DIR, '臺灣普惠科技應用痛點研究_說說學伴.docx')
    doc.save(docx_path)
    print(f'Word document saved: {docx_path}')
    return docx_path

if __name__ == '__main__':
    docx_path = build_document()
    
    # Convert to PDF
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        print(f'PDF document saved: {pdf_path}')
    except Exception as e:
        print(f'PDF conversion note: {e}')
        print('You can also convert to PDF by opening the .docx file in Word and using "Save As PDF".')
